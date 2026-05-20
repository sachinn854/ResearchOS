from docx import Document

from backend.ingestion.parsers.base import BaseParser, ParsedPage


class DOCXParser(BaseParser):
    def parse(self, source: str) -> list[ParsedPage]:
        doc = Document(source)
        sections: list[ParsedPage] = []
        current_heading: str | None = None
        current_paragraphs: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            if para.style.name.startswith("Heading"):
                # Flush the previous section before starting a new heading
                if current_paragraphs:
                    body = "\n".join(current_paragraphs)
                    combined = f"{current_heading}\n{body}" if current_heading else body
                    sections.append(ParsedPage(text=combined, page_number=None))
                current_heading = text
                current_paragraphs = []
            else:
                current_paragraphs.append(text)

        # Flush the last section
        if current_paragraphs:
            body = "\n".join(current_paragraphs)
            combined = f"{current_heading}\n{body}" if current_heading else body
            sections.append(ParsedPage(text=combined, page_number=None))

        # Fallback: no headings found — treat whole document as one page
        if not sections:
            full_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
            sections.append(ParsedPage(text=full_text, page_number=None))

        return sections
